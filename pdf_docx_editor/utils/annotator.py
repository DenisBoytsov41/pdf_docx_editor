import os
from PyPDF2 import PdfReader, PdfWriter  # Работа с PDF (чтение и запись)
from docx import Document  # Работа с DOCX-документами
import fitz  # PyMuPDF — работа с PDF на низком уровне (рендеринг, аннотации и т.д.)

def annotate_pdf(filepath):
    """
    Добавляет текстовую аннотацию к каждой странице PDF-файла.
    
    Аргументы:
        filepath (str): Путь к исходному PDF-файлу.
    
    Возвращает:
        str: Путь к аннотированному PDF-файлу.
    """
    # Открываем PDF-файл для обработки
    pdf_document = fitz.open(filepath)

    # Проходим по всем страницам в документе
    for page in pdf_document:
        # Создаём прямоугольную область для текста аннотации
        rect = fitz.Rect(100, 500, 200, 550)
        # Вставляем текст "Комментарий преподавателя" в указанную область
        page.insert_textbox(rect, "Комментарий преподавателя", fontsize=12, color=(1, 0, 0))  # Красный текст

    # Формируем путь для сохранения аннотированного файла
    annotated_filepath = filepath.replace('.pdf', '_annotated.pdf')
    # Сохраняем аннотированный файл
    pdf_document.save(annotated_filepath)
    # Закрываем файл
    pdf_document.close()

    # Возвращаем путь к аннотированному файлу
    return annotated_filepath

def annotate_docx(filepath):
    """
    Добавляет комментарии к абзацам в DOCX-документе, содержащим слово "ошибка".
    
    Аргументы:
        filepath (str): Путь к исходному DOCX-файлу.
    
    Возвращает:
        str: Путь к аннотированному DOCX-файлу.
    """
    # Открываем документ DOCX для обработки
    doc = Document(filepath)

    # Проходим по всем абзацам в документе
    for paragraph in doc.paragraphs:
        # Проверяем, содержит ли абзац слово "ошибка"
        if "ошибка" in paragraph.text:
            # Добавляем комментарий к абзацу (жирным шрифтом)
            paragraph.add_run(" (Комментарий: исправить форматирование)").bold = True

    # Формируем путь для сохранения аннотированного файла
    annotated_filepath = filepath.replace('.docx', '_annotated.docx')
    # Сохраняем аннотированный файл
    doc.save(annotated_filepath)

    # Возвращаем путь к аннотированному файлу
    return annotated_filepath
